"""
Engine implementations for Word COM and LibreOffice UNO.

Provides document lifecycle, selectors, and actions mapped to
native automation APIs.
"""

import logging
import platform
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Protocol

logger = logging.getLogger(__name__)


class EngineDocument(Protocol):
    """Protocol for engine-specific document handles."""

    pass


class EngineRange(Protocol):
    """Protocol for engine-specific range/selection objects."""

    pass


class Engine(ABC):
    """Abstract base for document automation engines."""

    @abstractmethod
    def open_document(self, path: Path) -> Any:
        """Open document and return engine-specific handle."""
        pass

    @abstractmethod
    def close_document(self, doc: Any) -> None:
        """Close document without saving."""
        pass

    @abstractmethod
    def save_as_new_docx(self, doc: Any, out_path: Path) -> None:
        """Save document to new .docx file."""
        pass

    @abstractmethod
    def export_pdf(self, doc: Any, out_pdf: Path) -> None:
        """Export document to PDF."""
        pass

    @abstractmethod
    def select_by_style(self, doc: Any, styles: list[str]) -> list[Any]:
        """Select ranges matching given style names."""
        pass

    @abstractmethod
    def select_by_regex(
        self, doc: Any, pattern: str, scope: str, flags: list[str], page_range: str | None
    ) -> list[Any]:
        """Select ranges matching regex pattern."""
        pass

    @abstractmethod
    def select_by_bookmark(self, doc: Any, names: list[str]) -> list[Any]:
        """Select bookmark ranges."""
        pass

    @abstractmethod
    def select_by_content_control(self, doc: Any, titles: list[str]) -> list[Any]:
        """Select content control ranges (Word only)."""
        pass

    @abstractmethod
    def select_by_table(
        self, doc: Any, index: int | None, style: str | None, contains_text: str | None
    ) -> list[Any]:
        """Select tables by criteria."""
        pass

    @abstractmethod
    def select_by_range(
        self,
        doc: Any,
        section: str | None,
        paragraph_indexes: list[int] | None,
        pages: str | None,
    ) -> list[Any]:
        """Select ranges by section/paragraph/page."""
        pass

    @abstractmethod
    def apply_paragraph_format(self, ranges: list[Any], fmt: dict[str, Any]) -> int:
        """Apply paragraph formatting to ranges."""
        pass

    @abstractmethod
    def apply_style(self, ranges: list[Any], style_name: str) -> int:
        """Apply named style to ranges."""
        pass

    @abstractmethod
    def apply_numbering(self, ranges: list[Any], numbering: dict[str, Any]) -> int:
        """Apply numbering/outline settings."""
        pass

    @abstractmethod
    def set_headers_footers(self, doc: Any, config: dict[str, Any]) -> None:
        """Set section headers and footers."""
        pass

    @abstractmethod
    def update_fields_and_toc(self, doc: Any, update_all: bool, update_toc: bool) -> None:
        """Update fields and table of contents."""
        pass

    @abstractmethod
    def find_replace(
        self,
        doc: Any,
        find: str,
        replace: str,
        regex: bool,
        wildcards: bool,
        whole_word: bool,
        match_case: bool,
    ) -> int:
        """Find and replace text."""
        pass

    @abstractmethod
    def apply_page_setup(self, doc: Any, setup: dict[str, Any]) -> None:
        """Apply page setup (margins, orientation, size)."""
        pass

    @abstractmethod
    def insert_section_break(
        self, doc: Any, before_selector: bool, break_type: str
    ) -> None:
        """Insert section break."""
        pass

    @abstractmethod
    def replace_bookmark_text(self, doc: Any, name: str, text: str) -> None:
        """Replace bookmark text content."""
        pass

    @abstractmethod
    def replace_content_control_text(self, doc: Any, title_or_tag: str, text: str) -> None:
        """Replace content control text (Word only)."""
        pass

    @abstractmethod
    def format_table(self, doc: Any, config: dict[str, Any]) -> None:
        """Format table properties and style."""
        pass

    @abstractmethod
    def insert_image(self, doc: Any, config: dict[str, Any]) -> None:
        """Insert image inline or floating."""
        pass

    @abstractmethod
    def raw_word_com(self, doc: Any, commands: list[dict[str, Any]]) -> None:
        """Execute raw Word COM commands (Word only)."""
        pass

    @abstractmethod
    def snapshot(self, doc: Any) -> dict[str, Any]:
        """Capture document snapshot for comparison."""
        pass

    @abstractmethod
    def shutdown(self) -> None:
        """Clean shutdown of engine."""
        pass


class WordComEngine(Engine):
    """Microsoft Word COM automation engine (Windows only)."""

    # Word constants
    wdFindStop = 0
    wdFindContinue = 1
    wdReplaceAll = 2
    wdWord = 2
    wdStory = 6
    wdCharacter = 1
    wdCell = 12
    wdAlertsNone = 0
    wdCollapseEnd = 0
    wdCollapseStart = 1
    wdParagraph = 4
    wdExtend = 1

    def __init__(self) -> None:
        """Initialize Word application via COM."""
        try:
            import win32com.client
            import win32com.client.gencache
            import pythoncom
            import sys
            import shutil
            import os
            from pathlib import Path
        except ImportError as e:
            raise RuntimeError(
                "pywin32 not available. Install with: pip install pywin32"
            ) from e

        # Force generation of static typelib
        logger.info("Initializing Word COM engine with forced type library generation")
        try:
            # Reset the gen_py directory
            gen_py = Path(win32com.client.gencache.GetGeneratePath())
            if gen_py.exists():
                shutil.rmtree(str(gen_py))
            os.makedirs(str(gen_py))
            
            # Reset the cache
            win32com.client.gencache.Rebuild()
            
            # Known Word type library info
            # Microsoft Word 16.0 Object Library (Office 2016+)
            win32com.client.gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 7)
            
            # Initialize Word with makepy support
            logger.info("Creating Word application instance")
            self.app = win32com.client.gencache.EnsureDispatch("Word.Application")
            self.app.Visible = False
            self.app.DisplayAlerts = self.wdAlertsNone
            
        except Exception as e:
            logger.error(f"COM initialization error: {str(e)}")
            # Try one more time with just basic Dispatch as fallback
            try:
                logger.info("Attempting fallback initialization")
                self.app = win32com.client.Dispatch("Word.Application")
                self.app.Visible = False
                self.app.DisplayAlerts = 0  # Can't use wdAlertsNone in dynamic dispatch
            except Exception as e2:
                raise RuntimeError(f"Failed to initialize Word COM (both attempts): {str(e)} / {str(e2)}") from e

    def open_document(self, path: Path) -> Any:
        """Open document via Word COM."""
        logger.debug(f"Opening document: {path}")
        doc = self.app.Documents.Open(str(path.absolute()))
        return doc

    def close_document(self, doc: Any) -> None:
        """Close document without saving."""
        try:
            doc.Close(SaveChanges=False)
        except Exception as e:
            logger.warning(f"Error closing document: {e}")

    def save_as_new_docx(self, doc: Any, out_path: Path) -> None:
        """Save as new .docx file."""
        logger.debug(f"Saving DOCX: {out_path}")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        # wdFormatXMLDocument = 12
        doc.SaveAs2(str(out_path.absolute()), FileFormat=12)

    def export_pdf(self, doc: Any, out_pdf: Path) -> None:
        """Export to PDF using Word's native exporter."""
        logger.info(f"Exporting PDF via Word COM: {out_pdf}")
        out_pdf.parent.mkdir(parents=True, exist_ok=True)
        # wdExportFormatPDF = 17
        doc.ExportAsFixedFormat(
            OutputFileName=str(out_pdf.absolute()),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,  # wdExportOptimizeForPrint
        )

    def select_by_style(self, doc: Any, styles: list[str]) -> list[Any]:
        """Select paragraphs by style."""
        ranges = []
        for para in doc.Paragraphs:
            if para.Style.NameLocal in styles:
                ranges.append(para.Range)
        logger.debug(f"Selected {len(ranges)} paragraphs by style: {styles}")
        return ranges

    def select_by_regex(
        self, doc: Any, pattern: str, scope: str, flags: list[str], page_range: str | None
    ) -> list[Any]:
        """Select ranges matching regex using Word Find."""
        ranges = []
        find = doc.Content.Find
        find.ClearFormatting()
        find.Text = pattern

        # Map flags to Word Find settings
        find.MatchWildcards = "WILDCARDS" in (flags or [])
        find.MatchCase = "MATCH_CASE" in (flags or [])
        find.MatchWholeWord = "WHOLE_WORD" in (flags or [])

        while find.Execute():
            ranges.append(doc.Content.Duplicate)
            # Move past current match
            doc.Content.Collapse(self.wdCollapseEnd)

        logger.debug(f"Selected {len(ranges)} ranges by regex: {pattern}")
        return ranges

    def select_by_bookmark(self, doc: Any, names: list[str]) -> list[Any]:
        """Select bookmark ranges."""
        ranges = []
        for name in names:
            if doc.Bookmarks.Exists(name):
                ranges.append(doc.Bookmarks(name).Range)
        logger.debug(f"Selected {len(ranges)} bookmarks")
        return ranges

    def select_by_content_control(self, doc: Any, titles: list[str]) -> list[Any]:
        """Select content controls by title/tag."""
        ranges = []
        for cc in doc.ContentControls:
            if cc.Title in titles or cc.Tag in titles:
                ranges.append(cc.Range)
        logger.debug(f"Selected {len(ranges)} content controls")
        return ranges

    def select_by_table(
        self, doc: Any, index: int | None, style: str | None, contains_text: str | None
    ) -> list[Any]:
        """Select tables by criteria."""
        ranges = []
        for i, table in enumerate(doc.Tables, start=1):
            if index is not None and i != index:
                continue
            if style is not None and table.Style.NameLocal != style:
                continue
            if contains_text is not None and contains_text not in table.Range.Text:
                continue
            ranges.append(table.Range)
        logger.debug(f"Selected {len(ranges)} tables")
        return ranges

    def select_by_range(
        self,
        doc: Any,
        section: str | None,
        paragraph_indexes: list[int] | None,
        pages: str | None,
    ) -> list[Any]:
        """Select by section/paragraph/page ranges."""
        ranges = []

        if section:
            if section == "all":
                for sec in doc.Sections:
                    ranges.append(sec.Range)
            else:
                sec_idx = int(section)
                if 1 <= sec_idx <= doc.Sections.Count:
                    ranges.append(doc.Sections(sec_idx).Range)

        if paragraph_indexes:
            for idx in paragraph_indexes:
                if 1 <= idx <= doc.Paragraphs.Count:
                    ranges.append(doc.Paragraphs(idx).Range)

        logger.debug(f"Selected {len(ranges)} ranges by range selector")
        return ranges

    def apply_paragraph_format(self, ranges: list[Any], fmt: dict[str, Any]) -> int:
        """Apply paragraph formatting to ranges."""
        count = 0
        for rng in ranges:
            pf = rng.ParagraphFormat
            if "line_spacing" in fmt:
                pf.LineSpacingRule = 5
                pf.LineSpacing = float(fmt["line_spacing"]) * 12.0  # multiple × 12pt
            if "space_before" in fmt:
                pf.SpaceBefore = self._parse_unit(fmt["space_before"])
            if "space_after" in fmt:
                pf.SpaceAfter = self._parse_unit(fmt["space_after"])
            if "keep_with_next" in fmt:
                pf.KeepWithNext = fmt["keep_with_next"]
            if "page_break_before" in fmt:
                pf.PageBreakBefore = fmt["page_break_before"]
            if "widow_control" in fmt:
                pf.WidowControl = fmt["widow_control"]
            if "first_line_indent" in fmt:
                pf.FirstLineIndent = self._parse_unit(fmt["first_line_indent"])
            if "left_indent" in fmt:
                pf.LeftIndent = self._parse_unit(fmt["left_indent"])
            if "right_indent" in fmt:
                pf.RightIndent = self._parse_unit(fmt["right_indent"])
            count += 1
        logger.debug(f"Applied paragraph format to {count} ranges")
        return count

    def apply_style(self, ranges: list[Any], style_name: str) -> int:
        """Apply named style to ranges."""
        count = 0
        for rng in ranges:
            rng.Style = style_name
            count += 1
        logger.debug(f"Applied style '{style_name}' to {count} ranges")
        return count

    def apply_numbering(self, ranges: list[Any], numbering: dict[str, Any]) -> int:
        """Apply numbering/outline settings."""
        count = 0
        for rng in ranges:
            if "outline_level" in numbering:
                rng.ParagraphFormat.OutlineLevel = numbering["outline_level"]
            if "list_template" in numbering:
                template_name = numbering["list_template"]
                # Find list template by name
                for lt in rng.ListTemplates:
                    if lt.Name == template_name:
                        rng.ListFormat.ApplyListTemplate(lt)
                        break
            if "restart_at" in numbering:
                restart = numbering["restart_at"]
                if restart != "auto":
                    rng.ListFormat.ListValue = int(restart)
            count += 1
        logger.debug(f"Applied numbering to {count} ranges")
        return count

    def set_headers_footers(self, doc: Any, config: dict[str, Any]) -> None:
        """Set section headers and footers."""
        for section in doc.Sections:
            if "header" in config:
                hdr_cfg = config["header"]
                # wdHeaderFooterPrimary = 1
                hdr = section.Headers(1)
                self._set_header_footer_text(hdr, hdr_cfg)

            if "footer" in config:
                ftr_cfg = config["footer"]
                # wdHeaderFooterPrimary = 1
                ftr = section.Footers(1)
                self._set_header_footer_text(ftr, ftr_cfg)

            if "different_first_page" in config:
                section.PageSetup.DifferentFirstPageHeaderFooter = config[
                    "different_first_page"
                ]
            if "different_odd_even" in config:
                section.PageSetup.OddAndEvenPagesHeaderFooter = config["different_odd_even"]

        logger.debug("Set headers/footers")

    def _set_header_footer_text(self, hf: Any, cfg: dict[str, str]) -> None:
        """Set header/footer text with alignment."""
        hf.Range.Text = ""
        if "left" in cfg:
            hf.Range.InsertAfter(cfg["left"])
        if "center" in cfg:
            hf.Range.ParagraphFormat.Alignment = 1  # wdAlignParagraphCenter
            hf.Range.InsertAfter("\t" + cfg["center"])
        if "right" in cfg:
            hf.Range.ParagraphFormat.Alignment = 2  # wdAlignParagraphRight
            hf.Range.InsertAfter("\t" + cfg["right"])

        # Replace field codes
        text = hf.Range.Text
        text = text.replace("{PAGE}", "").replace("{NUMPAGES}", "")
        hf.Range.Text = text

        # Insert actual fields
        if "{PAGE}" in str(cfg):
            hf.Range.Fields.Add(hf.Range, 33)  # wdFieldPage
        if "{NUMPAGES}" in str(cfg):
            hf.Range.Fields.Add(hf.Range, 26)  # wdFieldNumPages

    def update_fields_and_toc(self, doc: Any, update_all: bool, update_toc: bool) -> None:
        """Update fields and table of contents."""
        if update_all:
            doc.Fields.Update()
            logger.debug("Updated all fields")

        if update_toc:
            for toc in doc.TablesOfContents:
                toc.Update()
            logger.debug(f"Updated {doc.TablesOfContents.Count} TOCs")

    def find_replace(
        self,
        doc: Any,
        find: str,
        replace: str,
        regex: bool,
        wildcards: bool,
        whole_word: bool,
        match_case: bool,
    ) -> int:
        """Find and replace text."""
        find_obj = doc.Content.Find
        find_obj.ClearFormatting()
        find_obj.Replacement.ClearFormatting()

        find_obj.Text = find
        find_obj.Replacement.Text = replace
        find_obj.MatchWildcards = wildcards or regex
        find_obj.MatchWholeWord = whole_word
        find_obj.MatchCase = match_case

        # wdReplaceAll = 2
        count = find_obj.Execute(Replace=self.wdReplaceAll, Forward=True)
        logger.debug(f"Find/replace: {count} replacements")
        return count if count else 0

    def apply_page_setup(self, doc: Any, setup: dict[str, Any]) -> None:
        """Apply page setup to all sections."""
        for section in doc.Sections:
            ps = section.PageSetup

            if "margins" in setup:
                margins = setup["margins"]
                if "top" in margins:
                    ps.TopMargin = self._parse_unit(margins["top"])
                if "bottom" in margins:
                    ps.BottomMargin = self._parse_unit(margins["bottom"])
                if "left" in margins:
                    ps.LeftMargin = self._parse_unit(margins["left"])
                if "right" in margins:
                    ps.RightMargin = self._parse_unit(margins["right"])

            if "orientation" in setup:
                # wdOrientPortrait = 0, wdOrientLandscape = 1
                ps.Orientation = 0 if setup["orientation"] == "portrait" else 1

            if "paper_size" in setup:
                # Could map paper sizes here (A4, Letter, etc.)
                pass

        logger.debug("Applied page setup")

    def insert_section_break(
        self, doc: Any, before_selector: bool, break_type: str
    ) -> None:
        """Insert section break."""
        # wdSectionBreakNextPage = 2, wdSectionBreakContinuous = 3, etc.
        break_map = {
            "next_page": 2,
            "continuous": 3,
            "even_page": 4,
            "odd_page": 5,
        }
        break_code = break_map.get(break_type, 2)

        selection = doc.Application.Selection
        selection.InsertBreak(Type=break_code)
        logger.debug(f"Inserted section break: {break_type}")

    def replace_bookmark_text(self, doc: Any, name: str, text: str) -> None:
        """Replace bookmark text content."""
        if doc.Bookmarks.Exists(name):
            bm = doc.Bookmarks(name)
            bm.Range.Text = text
            logger.debug(f"Replaced bookmark '{name}'")

    def replace_content_control_text(self, doc: Any, title_or_tag: str, text: str) -> None:
        """Replace content control text."""
        for cc in doc.ContentControls:
            if cc.Title == title_or_tag or cc.Tag == title_or_tag:
                cc.Range.Text = text
                logger.debug(f"Replaced content control '{title_or_tag}'")

    def format_table(self, doc: Any, config: dict[str, Any]) -> None:
        """Format table properties."""
        index = config.get("index")
        if index and 1 <= index <= doc.Tables.Count:
            table = doc.Tables(index)

            if "style" in config:
                table.Style = config["style"]
            if "autofit" in config and config["autofit"]:
                table.AutoFitBehavior(1)  # wdAutoFitContent
            if "header_row" in config:
                table.Rows(1).HeadingFormat = config["header_row"]
            if "banded_rows" in config:
                table.ApplyStyleRowBands = config["banded_rows"]

            logger.debug(f"Formatted table {index}")

    def insert_image(self, doc: Any, config: dict[str, Any]) -> None:
        """Insert image inline or floating."""
        path = config["path"]
        anchor = config.get("anchor", "inline")

        selection = doc.Application.Selection
        if anchor == "inline":
            shape = selection.InlineShapes.AddPicture(FileName=path, LinkToFile=False)
        else:
            shape = selection.Shapes.AddPicture(FileName=path, LinkToFile=False)

        if "width" in config:
            shape.Width = self._parse_unit(config["width"])
        if "height" in config:
            shape.Height = self._parse_unit(config["height"])

        logger.debug(f"Inserted image: {path}")

    def raw_word_com(self, doc: Any, commands: list[dict[str, Any]]) -> None:
        """Execute raw Word COM commands (advanced escape hatch)."""
        logger.warning("Executing raw_word_com commands (non-portable)")
        for cmd in commands:
            target = cmd["target"]
            prop = cmd["property"]
            value = cmd["value"]

            try:
                # Evaluate target path (e.g., "Selection", "Sections[1]")
                obj = eval(f"doc.{target}")
                # Set property
                parts = prop.split(".")
                for part in parts[:-1]:
                    obj = getattr(obj, part)
                setattr(obj, parts[-1], value)
                logger.debug(f"Raw COM: {target}.{prop} = {value}")
            except Exception as e:
                logger.error(f"Raw COM command failed: {e}")

    def snapshot(self, doc: Any) -> dict[str, Any]:
        """Capture document snapshot."""
        snap = {
            "paragraph_count": doc.Paragraphs.Count,
            "bookmark_count": doc.Bookmarks.Count,
            "inline_shape_count": doc.InlineShapes.Count,
            "content_control_count": doc.ContentControls.Count,
            "tables_count": doc.Tables.Count,
            "headings_by_level": {},
        }

        # Count headings by level
        for para in doc.Paragraphs:
            style = para.Style.NameLocal
            if "Heading" in style:
                snap["headings_by_level"][style] = snap["headings_by_level"].get(style, 0) + 1

        return snap

    def shutdown(self) -> None:
        """Shutdown Word application."""
        try:
            self.app.Quit()
            logger.info("Word COM engine shutdown")
        except Exception as e:
            logger.warning(f"Error during shutdown: {e}")

    @staticmethod
    def _parse_unit(value: str) -> float:
        """Parse unit string to points (Word's base unit)."""
        if isinstance(value, (int, float)):
            return float(value)

        value = str(value).strip().lower()
        if value.endswith("pt"):
            return float(value[:-2])
        elif value.endswith("cm"):
            return float(value[:-2]) * 28.35
        elif value.endswith("in"):
            return float(value[:-2]) * 72
        elif value.endswith("mm"):
            return float(value[:-2]) * 2.835
        else:
            return float(value)


class LibreUnoEngine(Engine):
    """LibreOffice UNO automation engine (limited feature support)."""

    def __init__(self) -> None:
        """Initialize LibreOffice UNO connection."""
        logger.warning("Using LibreOffice UNO engine (limited capabilities)")
        # TODO: Implement UNO connection
        raise NotImplementedError("LibreOffice UNO engine not yet implemented")


class DocxEngine(Engine):
    """Engine implementation using python-docx for cross-platform operation.

    This provides a subset of the Word COM features implemented using
    the python-docx library. Some advanced features (bookmarks, content
    controls, updating TOC/fields, full section break control, raw COM)
    are intentionally left as NotImplemented or limited because python-docx
    and the docx XML model do not expose them directly.
    """

    def __init__(self) -> None:
        # import locally so module import does not fail when python-docx
        # isn't installed; __init__ will raise then and caller can install
        try:
            from docx import Document as DocxDocument  # type: ignore
            from docx.shared import Pt, Cm, Inches  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "python-docx is required for DocxEngine. Install with: pip install python-docx"
            ) from e

        self._DocxDocument = DocxDocument
        self._Pt = Pt
        self._Cm = Cm
        self._Inches = Inches
        self.doc = None

    def open_document(self, path: Path) -> Any:
        """Open a .docx file and return the python-docx Document object."""
        self.doc = self._DocxDocument(str(path))
        return self.doc

    def close_document(self, doc: Any) -> None:
        """python-docx has no close; just drop reference."""
        self.doc = None

    def save_as_new_docx(self, doc: Any, out_path: Path) -> None:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))

    def export_pdf(self, doc: Any, out_pdf: Path) -> None:
        """Export to PDF is not provided by python-docx; recommend using
        unoconv/libreoffice in a separate step for reliable PDF output.
        """
        raise NotImplementedError(
            "Export to PDF is not supported by DocxEngine; run LibreOffice conversion separately."
        )

    def select_by_style(self, doc: Any, styles: list[str]) -> list[Any]:
        ranges = []
        for para in doc.paragraphs:
            try:
                name = para.style.name if para.style is not None else None
            except Exception:
                name = None
            if name in styles:
                ranges.append(para)
        logger.debug(f"Selected {len(ranges)} paragraphs by style: {styles}")
        return ranges

    def select_by_regex(
        self, doc: Any, pattern: str, scope: str, flags: list[str], page_range: str | None
    ) -> list[Any]:
        ranges = []
        re_flags = 0
        if flags and "IGNORECASE" in flags:
            re_flags |= re.I
        compiled = re.compile(pattern, re_flags)

        # Search paragraphs and table cells
        for para in doc.paragraphs:
            if compiled.search(para.text or ""):
                ranges.append(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if compiled.search(cell.text or ""):
                        # append cell object
                        ranges.append(cell)

        logger.debug(f"Selected {len(ranges)} ranges by regex: {pattern}")
        return ranges

    def select_by_bookmark(self, doc: Any, names: list[str]) -> list[Any]:
        # python-docx does not have bookmark API; return empty list
        logger.debug("Bookmarks not supported by python-docx")
        return []

    def select_by_content_control(self, doc: Any, titles: list[str]) -> list[Any]:
        logger.debug("Content controls not supported by python-docx")
        return []

    def select_by_table(
        self, doc: Any, index: int | None, style: str | None, contains_text: str | None
    ) -> list[Any]:
        ranges = []
        for i, table in enumerate(doc.tables, start=1):
            if index is not None and i != index:
                continue
            if style is not None:
                try:
                    if table.style.name != style:
                        continue
                except Exception:
                    pass
            if contains_text is not None and contains_text not in table.cell(0, 0).text:
                # crude contains test; fall back to full table text
                full = "\n".join([c.text for row in table.rows for c in row.cells])
                if contains_text not in full:
                    continue
            ranges.append(table)
        logger.debug(f"Selected {len(ranges)} tables")
        return ranges

    def select_by_range(
        self,
        doc: Any,
        section: str | None,
        paragraph_indexes: list[int] | None,
        pages: str | None,
    ) -> list[Any]:
        ranges = []
        if paragraph_indexes:
            for idx in paragraph_indexes:
                if 1 <= idx <= len(doc.paragraphs):
                    ranges.append(doc.paragraphs[idx - 1])
        # section/pages selectors are not fully supported
        logger.debug(f"Selected {len(ranges)} ranges by range selector")
        return ranges

    def apply_paragraph_format(self, ranges: list[Any], fmt: dict[str, Any]) -> int:
        count = 0
        for para in ranges:
            pf = para.paragraph_format
            if "line_spacing" in fmt:
                try:
                    pf.line_spacing = float(fmt["line_spacing"])  # simple numeric multiplier
                except Exception:
                    pass
            if "space_before" in fmt:
                pf.space_before = self._parse_unit(fmt["space_before"])
            if "space_after" in fmt:
                pf.space_after = self._parse_unit(fmt["space_after"])
            if "first_line_indent" in fmt:
                pf.first_line_indent = self._parse_unit(fmt["first_line_indent"])
            if "left_indent" in fmt:
                pf.left_indent = self._parse_unit(fmt["left_indent"])
            if "right_indent" in fmt:
                pf.right_indent = self._parse_unit(fmt["right_indent"])
            count += 1
        logger.debug(f"Applied paragraph format to {count} ranges")
        return count

    def apply_style(self, ranges: list[Any], style_name: str) -> int:
        count = 0
        for para in ranges:
            try:
                para.style = style_name
                count += 1
            except Exception:
                logger.debug(f"Failed to set style '{style_name}' on paragraph")
        logger.debug(f"Applied style '{style_name}' to {count} ranges")
        return count

    def apply_numbering(self, ranges: list[Any], numbering: dict[str, Any]) -> int:
        # python-docx has only limited support for numbering via styles
        count = 0
        list_style = numbering.get("list_style")
        for para in ranges:
            if list_style:
                try:
                    para.style = list_style
                except Exception:
                    pass
            count += 1
        logger.debug(f"Applied numbering (limited) to {count} ranges")
        return count

    def set_headers_footers(self, doc: Any, config: dict[str, Any]) -> None:
        for section in doc.sections:
            if "header" in config:
                hdr = section.header
                hdr.is_linked_to_previous = False
                hdr_para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
                hdr_para.text = config["header"].get("center", "")
            if "footer" in config:
                ftr = section.footer
                ftr_para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
                ftr_para.text = config["footer"].get("center", "")
        logger.debug("Set headers/footers (limited)")

    def update_fields_and_toc(self, doc: Any, update_all: bool, update_toc: bool) -> None:
        raise NotImplementedError("python-docx cannot update fields or TOC programmatically")

    def find_replace(
        self,
        doc: Any,
        find: str,
        replace: str,
        regex: bool,
        wildcards: bool,
        whole_word: bool,
        match_case: bool,
    ) -> int:
        count = 0
        if regex:
            pattern = re.compile(find)
            for para in doc.paragraphs:
                new = pattern.sub(replace, para.text)
                if new != para.text:
                    para.text = new
                    count += 1
        else:
            for para in doc.paragraphs:
                if (match_case and find in para.text) or (not match_case and find.lower() in para.text.lower()):
                    if match_case:
                        para.text = para.text.replace(find, replace)
                    else:
                        # case-insensitive replace
                        para.text = re.compile(re.escape(find), re.I).sub(replace, para.text)
                    count += 1
        logger.debug(f"Find/replace: {count} paragraphs modified")
        return count

    def apply_page_setup(self, doc: Any, setup: dict[str, Any]) -> None:
        for section in doc.sections:
            if "margins" in setup:
                margins = setup["margins"]
                if "top" in margins:
                    section.top_margin = self._parse_unit(margins["top"])
                if "bottom" in margins:
                    section.bottom_margin = self._parse_unit(margins["bottom"])
                if "left" in margins:
                    section.left_margin = self._parse_unit(margins["left"])
                if "right" in margins:
                    section.right_margin = self._parse_unit(margins["right"])
            if "orientation" in setup:
                # python-docx uses section.orientation but mapping specifics may vary
                pass
        logger.debug("Applied page setup (limited)")

    def insert_section_break(
        self, doc: Any, before_selector: bool, break_type: str
    ) -> None:
        # python-docx support for inserting specific section breaks is limited.
        try:
            from docx.enum.section import WD_SECTION  # type: ignore

            mapping = {
                "next_page": WD_SECTION.NEW_PAGE,
                "continuous": WD_SECTION.CONTINUOUS,
                "even_page": WD_SECTION.EVEN_PAGE,
                "odd_page": WD_SECTION.ODD_PAGE,
            }
            kind = mapping.get(break_type, WD_SECTION.NEW_PAGE)
            doc.add_section(kind)
        except Exception:
            raise NotImplementedError("Insert section break is limited in python-docx")

    def replace_bookmark_text(self, doc: Any, name: str, text: str) -> None:
        # Not supported
        logger.debug("replace_bookmark_text not supported in DocxEngine")

    def replace_content_control_text(self, doc: Any, title_or_tag: str, text: str) -> None:
        logger.debug("replace_content_control_text not supported in DocxEngine")

    def format_table(self, doc: Any, config: dict[str, Any]) -> None:
        index = config.get("index")
        if index and 1 <= index <= len(doc.tables):
            table = doc.tables[index - 1]
            if "style" in config:
                try:
                    table.style = config["style"]
                except Exception:
                    pass

    def insert_image(self, doc: Any, config: dict[str, Any]) -> None:
        path = config["path"]
        anchor = config.get("anchor", "inline")
        width = config.get("width")
        height = config.get("height")
        # Try to insert into document at end; more advanced placement requires run.add_picture
        try:
            if width:
                doc.add_picture(path, width=self._parse_unit(width))
            else:
                doc.add_picture(path)
        except Exception:
            logger.debug("Failed to add picture via document API")

    def raw_word_com(self, doc: Any, commands: list[dict[str, Any]]) -> None:
        logger.warning("raw_word_com not supported in DocxEngine (Word-only feature)")

    def snapshot(self, doc: Any) -> dict[str, Any]:
        snap = {
            "paragraph_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections),
        }
        return snap

    def shutdown(self) -> None:
        pass

    def _parse_unit(self, value: str):
        """Parse unit strings into python-docx length objects (Pt or Inches).

        Returns a float (points) for assignment to margin/indent properties which
        python-docx accepts as EMU-friendly floats (Pt(...) objects are accepted).
        """
        if isinstance(value, (int, float)):
            return self._Pt(float(value))

        value = str(value).strip().lower()
        if value.endswith("pt"):
            return self._Pt(float(value[:-2]))
        if value.endswith("cm"):
            return self._Cm(float(value[:-2]))
        if value.endswith("in"):
            return self._Inches(float(value[:-2]))
        if value.endswith("mm"):
            # convert mm to cm
            return self._Cm(float(value[:-2]) / 10.0)
        # fallback assume points
        try:
            return self._Pt(float(value))
        except Exception:
            return self._Pt(0)


    def open_document(self, path: Path) -> Any:
        raise NotImplementedError()

    def close_document(self, doc: Any) -> None:
        raise NotImplementedError()

    def save_as_new_docx(self, doc: Any, out_path: Path) -> None:
        raise NotImplementedError()

    def export_pdf(self, doc: Any, out_pdf: Path) -> None:
        """Export PDF via LibreOffice (fallback)."""
        logger.warning("PDF export via UNO may have different fidelity than Word")
        raise NotImplementedError()

    def select_by_style(self, doc: Any, styles: list[str]) -> list[Any]:
        raise NotImplementedError()

    def select_by_regex(
        self, doc: Any, pattern: str, scope: str, flags: list[str], page_range: str | None
    ) -> list[Any]:
        raise NotImplementedError()

    def select_by_bookmark(self, doc: Any, names: list[str]) -> list[Any]:
        raise NotImplementedError()

    def select_by_content_control(self, doc: Any, titles: list[str]) -> list[Any]:
        logger.warning("Content controls not supported in UNO engine")
        return []

    def select_by_table(
        self, doc: Any, index: int | None, style: str | None, contains_text: str | None
    ) -> list[Any]:
        raise NotImplementedError()

    def select_by_range(
        self,
        doc: Any,
        section: str | None,
        paragraph_indexes: list[int] | None,
        pages: str | None,
    ) -> list[Any]:
        raise NotImplementedError()

    def apply_paragraph_format(self, ranges: list[Any], fmt: dict[str, Any]) -> int:
        raise NotImplementedError()

    def apply_style(self, ranges: list[Any], style_name: str) -> int:
        raise NotImplementedError()

    def apply_numbering(self, ranges: list[Any], numbering: dict[str, Any]) -> int:
        logger.warning("Advanced numbering limited in UNO engine")
        return 0

    def set_headers_footers(self, doc: Any, config: dict[str, Any]) -> None:
        raise NotImplementedError()

    def update_fields_and_toc(self, doc: Any, update_all: bool, update_toc: bool) -> None:
        raise NotImplementedError()

    def find_replace(
        self,
        doc: Any,
        find: str,
        replace: str,
        regex: bool,
        wildcards: bool,
        whole_word: bool,
        match_case: bool,
    ) -> int:
        raise NotImplementedError()

    def apply_page_setup(self, doc: Any, setup: dict[str, Any]) -> None:
        raise NotImplementedError()

    def insert_section_break(
        self, doc: Any, before_selector: bool, break_type: str
    ) -> None:
        raise NotImplementedError()

    def replace_bookmark_text(self, doc: Any, name: str, text: str) -> None:
        raise NotImplementedError()

    def replace_content_control_text(self, doc: Any, title_or_tag: str, text: str) -> None:
        logger.warning("Content controls not supported in UNO engine")

    def format_table(self, doc: Any, config: dict[str, Any]) -> None:
        raise NotImplementedError()

    def insert_image(self, doc: Any, config: dict[str, Any]) -> None:
        raise NotImplementedError()

    def raw_word_com(self, doc: Any, commands: list[dict[str, Any]]) -> None:
        logger.warning("raw_word_com not supported in UNO engine (Word-only feature)")

    def snapshot(self, doc: Any) -> dict[str, Any]:
        raise NotImplementedError()

    def shutdown(self) -> None:
        pass


def pick_engine(engine_hint: str) -> Engine:
    """
    Select and initialize appropriate engine.

    Args:
        engine_hint: "auto", "word", or "libre"

    Returns:
        Initialized engine instance

    Raises:
        RuntimeError: If required engine unavailable
    """
    if engine_hint == "word":
        if platform.system() != "Windows":
            raise RuntimeError("Word COM engine requires Windows")
        return WordComEngine()

    elif engine_hint == "libre":
        return LibreUnoEngine()

    elif engine_hint == "docx":
        # Use python-docx engine (cross-platform)
        return DocxEngine()

    else:  # auto
        if platform.system() == "Windows":
            try:
                return WordComEngine()
            except Exception as e:
                logger.warning(f"Word COM unavailable: {e}")

        # Try python-docx as a cross-platform fallback first
        try:
            return DocxEngine()
        except Exception as e:
            logger.warning(f"DocxEngine unavailable: {e}")

        # Try UNO fallback
        try:
            return LibreUnoEngine()
        except NotImplementedError:
            raise RuntimeError(
                "No engine available. Install python-docx, Microsoft Word (Windows) "
                "or LibreOffice with UNO bindings."
            )